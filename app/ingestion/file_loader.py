from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


@dataclass
class RawDocument:
    doc_id: str
    source: str
    text: str
    source_type: str = "file"


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    lines = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(lines)


def read_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return _read_txt(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    raise ValueError(f"Unsupported file type: {path}")


def load_documents_from_directory(directory: str) -> list[RawDocument]:
    root = Path(directory)
    if not root.exists() or not root.is_dir():
        raise ValueError(f"Directory does not exist: {directory}")

    docs: list[RawDocument] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        text = read_file(path).strip()
        if not text:
            continue

        doc_id = str(path.resolve())
        docs.append(
            RawDocument(
                doc_id=doc_id,
                source=str(path.resolve()),
                text=text,
            )
        )

    return docs
